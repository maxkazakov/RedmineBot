from docx import Document
from docx.shared import Cm, Inches, Pt, Emu

head_names = ['№ п/п', 'Объект тестирования', 'Сценарий выполнения', 'Ожидаемый результат', 'Инициатор', 'Итог тестирования']
# cm
column_width = [1, 5, 7, 6, 3, 5]

def FillRowData(row, data, isHead = False):
    for i, text in enumerate(data):
        run = row[i].add_paragraph().add_run(text)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(10)
        row[i].width = Cm(column_width[i]).emu
        if isHead:
            run.bold = True

def BuildDocx(data):
    document = Document('template.docx')

    # заголовок
    paragraph = document.add_paragraph()
    run = paragraph.add_run('ТЕСТ-ПЛАН. ОБНОВЛЕНИЕ 4.1.9')
    run.style = 'Заголовок 1 Знак'
    paragraph.paragraph_format.first_line_indent = Cm(3.0)

    # цикл по категориям
    idx = 1
    for ctg, issues in data.items():
        # вписываем категорию
        paragraph = document.add_paragraph('Таблица {} {}'.format(idx, ctg))
        # paragraph.paragraph_format.space_after = Pt(12)
        paragraph.style = 'Caption'

        # строим таблицу
        table = document.add_table(rows=len(issues) + 1, cols=6)

        table.style = 'Table Grid'
        row_id = 0
        row = table.rows[row_id].cells
        # заголовок
        FillRowData(row, head_names, True)
        # данные
        for iss in issues:
            row_id += 1
            row = table.rows[row_id].cells
            iss.insert(0, '{}.'.format(row_id))
            FillRowData(row, iss)
        idx+=1

        # Че не работает то ??? Повесил width  в итоге на cell-ы, а не на column-ы(
        # for i, col in enumerate(table.columns):
        #     col.width = Cm(column_width[i]).emu

    document.save('Протокол тестирования.docx')

# BuildDocx(data)
